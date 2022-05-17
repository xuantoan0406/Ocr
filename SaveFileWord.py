from docx import Document


class SaveFileWord:
	def __init__(self):
		self.document = Document()
	
	def create_document(self,fileName):
		self.document.add_paragraph("hello word \t \t aaaaaaaa")
		self.document.add_paragraph("hello word \t \t aaaaaaaa")
		self.document.save(fileName)


a=SaveFileWord()
a.create_document("test.docx")