# -*- coding: utf-8 -*-
import cv2
import imutils
from docx import Document
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def show_img(img, sizeH):
	cv2.imshow("a", imutils.resize(img, height=sizeH))
	cv2.waitKey()


def write_text(listText, nameOut):
	with open(f'{nameOut}.txt', 'w', encoding="utf-8") as f:
		for textPage in listText:
			for text in textPage:
				f.writelines(text + '\n')
			f.writelines(3 * '\n')


def write_docx(listText, nameOut):
	document = Document()
	# style = document.styles['Normal']
	# font = style.font
	# font.name = 'Times New Roman'
	# font.size = Pt(11)
	
	for i, textPage in enumerate(listText):
		
		for text in textPage:
			document.add_paragraph(text)
		# document.add_page_number(document.sections[0].footer.paragraphs[0].add_run())
		# document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
		if i < len(listText) - 1:
			document.add_page_break()
	# p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	document.save(f"{nameOut}.docx")
